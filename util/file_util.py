import io
import os
import shutil
from dataclasses import dataclass
from datetime import timedelta, datetime

import fitz
from PIL.Image import Resampling
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Pt, RGBColor, Cm
from PIL import Image
from loguru import logger

from path import get_work_path


@dataclass
class DocData:
    title: str
    author: str
    institution: str
    doi: str
    desc: str
    img: str


def resize_image_if_needed(image_data, image_type, max_resolution=(2560, 1440)):
    """Resize image if its dimensions exceed the specified max resolution."""
    with Image.open(io.BytesIO(image_data)) as img:
        # Check if the image exceeds the max resolution
        if img.width > max_resolution[0] or img.height > max_resolution[1]:
            # Calculate the new size while maintaining the aspect ratio
            img.thumbnail(max_resolution, Resampling.BILINEAR)
            output = io.BytesIO()
            img.save(output, format=image_type)
            return output.getvalue()
        else:
            return image_data


def recover_pix(doc: Document, item):
    xref = item[0]  # xref of PDF image
    smask = item[1]  # xref of its /SMask

    # special case: /SMask or /Mask exists
    if smask > 0:
        pix0 = fitz.Pixmap(doc.extract_image(xref)["image"])
        if pix0.alpha:  # catch irregular situation
            pix0 = fitz.Pixmap(pix0, 0)  # remove alpha channel
        mask = fitz.Pixmap(doc.extract_image(smask)["image"])

        try:
            pix = fitz.Pixmap(pix0, mask)
        except:  # fallback to original base image in case of problems
            pix = fitz.Pixmap(doc.extract_image(xref)["image"])

        if pix0.n > 3:
            ext = "pam"
        else:
            ext = "png"

        return {  # create dictionary expected by caller
            "ext": ext,
            "colorspace": pix.colorspace.n,
            "image": pix.tobytes(ext),
        }

    # special case: /ColorSpace definition exists
    # to be sure, we convert these cases to RGB PNG images
    if "/ColorSpace" in doc.xref_object(xref, compressed=True):
        pix = fitz.Pixmap(doc, xref)
        pix = fitz.Pixmap(fitz.csRGB, pix)
        return {  # create dictionary expected by caller
            "ext": "png",
            "colorspace": 3,
            "image": pix.tobytes("png"),
        }
    return doc.extract_image(xref)


def get_image(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)

    page_count = doc.page_count  # number of pages

    xref_list = []
    img_list = []
    for pno in range(page_count):
        if len(img_list) >= 4:
            break

        il = doc.get_page_images(pno)
        for img in il:
            xref = img[0]
            if xref in xref_list:
                continue

            image = recover_pix(doc, img)
            imgdata = image["image"]

            imgdata = resize_image_if_needed(imgdata, image['ext'])

            img_file = os.path.join(
                os.path.dirname(pdf_path),
                f"page_{pno}_img_{xref}.{image['ext']}"
            )
            os.makedirs(os.path.dirname(img_file), exist_ok=True)
            with open(img_file, "wb") as fout:
                fout.write(imgdata)

            img_list.append(img_file)
            xref_list.append(xref)

    return img_list[0] if len(img_list) > 0 else ""


def write_to_docx(paper_list: list[DocData], output_file: str | bytes):
    document = Document()

    representation = document.add_paragraph().add_run(
        "本文内容为生成式AI对文章进行总结后得到，版权归原文作者所有。总结内容可靠性无保障，请仔细鉴别并以原文为准。")
    representation.font.size = Pt(11.5)
    representation.font.color.rgb = RGBColor(123, 125, 125)

    for data in paper_list:
        title_run = document.add_paragraph().add_run(f"\n\n{data.title}")
        title_run.font.size = Pt(16)
        title_run.bold = True

        p1 = document.add_paragraph()
        p1.paragraph_format.space_after = 5
        p1.paragraph_format.line_spacing = 1
        author_run = p1.add_run(data.author)
        author_run.font.size = Pt(12)
        author_run.font.color.rgb = RGBColor(123, 125, 125)

        p2 = document.add_paragraph()
        p2.paragraph_format.space_before = 5
        p2.paragraph_format.line_spacing = 1
        institution_run = p2.add_run(data.institution)
        institution_run.font.size = Pt(12)
        institution_run.font.color.rgb = RGBColor(123, 125, 125)

        url_run = document.add_paragraph().add_run(f"https://doi.org/{data.doi}")
        url_run.font.size = Pt(12)
        url_run.font.italic = True

        p3 = document.add_paragraph()
        p3.paragraph_format.line_spacing = 1.75
        desc_run = p3.add_run(data.desc)
        desc_run.font.size = Pt(13)

        if data.img != "":
            try:
                document.add_picture(data.img, width=Cm(13))
            except UnrecognizedImageError as e:
                logger.error(f'"{e}", {data.img}')

    document.save(output_file)


def compress_folder(yesterday: str):
    output_path = get_work_path()
    shutil.make_archive(
        f'{yesterday}-summary',
        'zip',
        output_path,
        f'{yesterday}-summary'
    )


def main() -> None:
    write_to_docx([
        DocData(
            "title",
            "Aye",
            "CAS",
            "12345",
            "desccccccccc",
            ""
        ),
        DocData(
            "title",
            "Aye",
            "CAS",
            "12345",
            "desccccccccc",
            ""
        )
    ], "test.docx")


if __name__ == '__main__':
    main()
